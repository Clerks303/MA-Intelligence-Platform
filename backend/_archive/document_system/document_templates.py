"""
Système de templates et génération automatique de documents
US-012: Templates intelligents et génération automatisée pour documents M&A

Ce module fournit:
- Templates configurables pour tous types de documents M&A
- Génération automatique à partir de données entreprises
- Moteur de rendu avec variables et logique conditionnelle
- Bibliothèque de templates standard (NDA, LOI, etc.)
- Personnalisation et branding automatique
- Export multi-formats (PDF, DOCX, HTML)
"""

import asyncio
import os
import json
import re
import uuid
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Union, Callable
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
import tempfile
import zipfile
from jinja2 import Environment, FileSystemLoader, Template, select_autoescape
from jinja2.exceptions import TemplateError

# Imports pour génération de documents
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

# Pour génération DOCX
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.logging_system import get_logger, LogCategory
from app.core.cache_manager import get_cache_manager

logger = get_logger("document_templates", LogCategory.DOCUMENT)


class TemplateType(str, Enum):
    """Types de templates disponibles"""
    NDA = "nda"
    LOI = "loi"  # Letter of Intent
    TERM_SHEET = "term_sheet"
    DUE_DILIGENCE = "due_diligence"
    FINANCIAL_REPORT = "financial_report"
    LEGAL_OPINION = "legal_opinion"
    VALUATION_REPORT = "valuation_report"
    MERGER_AGREEMENT = "merger_agreement"
    DISCLOSURE_LETTER = "disclosure_letter"
    BOARD_RESOLUTION = "board_resolution"
    SHAREHOLDER_AGREEMENT = "shareholder_agreement"
    EMPLOYMENT_AGREEMENT = "employment_agreement"
    CONSULTING_AGREEMENT = "consulting_agreement"
    CUSTOM = "custom"


class OutputFormat(str, Enum):
    """Formats de sortie supportés"""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    TXT = "txt"
    JSON = "json"


class TemplateEngine(str, Enum):
    """Moteurs de templates supportés"""
    JINJA2 = "jinja2"
    CUSTOM = "custom"


@dataclass
class TemplateVariable:
    """Variable de template"""
    name: str
    type: str  # string, number, date, boolean, list, dict
    description: str
    required: bool = False
    default_value: Any = None
    validation_pattern: Optional[str] = None
    choices: Optional[List[str]] = None


@dataclass
class TemplateMetadata:
    """Métadonnées d'un template"""
    
    # Identifiants
    template_id: str
    name: str
    description: str
    template_type: TemplateType
    
    # Versioning
    version: str = "1.0.0"
    author: str = ""
    created_at: datetime = field(default_factory=datetime.now)
    updated_at: datetime = field(default_factory=datetime.now)
    
    # Configuration
    engine: TemplateEngine = TemplateEngine.JINJA2
    supported_formats: List[OutputFormat] = field(default_factory=lambda: [OutputFormat.PDF, OutputFormat.DOCX])
    
    # Variables et validation
    variables: List[TemplateVariable] = field(default_factory=list)
    required_data_sources: List[str] = field(default_factory=list)
    
    # Métadonnées business
    legal_jurisdiction: Optional[str] = None
    applicable_regulations: List[str] = field(default_factory=list)
    confidentiality_level: str = "confidential"
    
    # Utilisation
    usage_count: int = 0
    rating: float = 0.0
    tags: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convertit en dictionnaire"""
        data = {}
        for field_name, field_def in self.__dataclass_fields__.items():
            value = getattr(self, field_name)
            
            if isinstance(value, datetime):
                data[field_name] = value.isoformat()
            elif isinstance(value, Enum):
                data[field_name] = value.value
            elif isinstance(value, list) and value and hasattr(value[0], '__dict__'):
                data[field_name] = [item.__dict__ if hasattr(item, '__dict__') else item for item in value]
            else:
                data[field_name] = value
        
        return data
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'TemplateMetadata':
        """Crée depuis un dictionnaire"""
        
        # Convertir les dates
        for date_field in ['created_at', 'updated_at']:
            if data.get(date_field) and isinstance(data[date_field], str):
                data[date_field] = datetime.fromisoformat(data[date_field])
        
        # Convertir les enums
        if 'template_type' in data:
            data['template_type'] = TemplateType(data['template_type'])
        if 'engine' in data:
            data['engine'] = TemplateEngine(data['engine'])
        if 'supported_formats' in data:
            data['supported_formats'] = [OutputFormat(fmt) for fmt in data['supported_formats']]
        
        # Convertir les variables
        if 'variables' in data and isinstance(data['variables'], list):
            variables = []
            for var_data in data['variables']:
                if isinstance(var_data, dict):
                    variables.append(TemplateVariable(**var_data))
                else:
                    variables.append(var_data)
            data['variables'] = variables
        
        return cls(**data)


@dataclass
class GenerationRequest:
    """Requête de génération de document"""
    
    template_id: str
    output_format: OutputFormat
    variables: Dict[str, Any] = field(default_factory=dict)
    
    # Métadonnées de génération
    generated_by: str = ""
    company_id: Optional[str] = None
    deal_id: Optional[str] = None
    
    # Options de personnalisation
    branding: Dict[str, Any] = field(default_factory=dict)
    custom_styles: Dict[str, Any] = field(default_factory=dict)
    
    # Options de sécurité
    watermark: Optional[str] = None
    password_protect: bool = False
    password: Optional[str] = None


@dataclass
class GenerationResult:
    """Résultat de génération de document"""
    
    success: bool
    file_data: Optional[bytes] = None
    filename: str = ""
    mime_type: str = ""
    
    # Métadonnées de génération
    template_used: str = ""
    generation_time: float = 0.0
    generated_at: datetime = field(default_factory=datetime.now)
    
    # Erreurs
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    
    # Statistiques
    pages_generated: int = 0
    word_count: int = 0


class TemplateEngine_:
    """Moteur de rendu de templates"""
    
    def __init__(self, templates_dir: str = "templates"):
        self.templates_dir = Path(templates_dir)
        self.templates_dir.mkdir(exist_ok=True)
        
        # Moteur Jinja2
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=select_autoescape(['html', 'xml']),
            trim_blocks=True,
            lstrip_blocks=True
        )
        
        # Fonctions personnalisées pour templates
        self._setup_template_functions()
        
        # Cache des templates compilés
        self.template_cache = {}
    
    def _setup_template_functions(self):
        """Configure les fonctions personnalisées pour templates"""
        
        def format_currency(amount, currency="EUR"):
            """Formate un montant en devise"""
            if isinstance(amount, (int, float)):
                return f"{amount:,.2f} {currency}"
            return str(amount)
        
        def format_date(date_obj, format_str="%d/%m/%Y"):
            """Formate une date"""
            if isinstance(date_obj, datetime):
                return date_obj.strftime(format_str)
            elif isinstance(date_obj, str):
                try:
                    return datetime.fromisoformat(date_obj).strftime(format_str)
                except:
                    return date_obj
            return str(date_obj)
        
        def format_percentage(value, decimals=2):
            """Formate un pourcentage"""
            if isinstance(value, (int, float)):
                return f"{value:.{decimals}f}%"
            return str(value)
        
        def format_number(value, thousands_sep=" "):
            """Formate un nombre avec séparateurs"""
            if isinstance(value, (int, float)):
                return f"{value:,}".replace(",", thousands_sep)
            return str(value)
        
        def upper_first(text):
            """Met la première lettre en majuscule"""
            return text.capitalize() if isinstance(text, str) else str(text)
        
        def legal_entity_format(entity_name, legal_form):
            """Formate le nom d'une entité légale"""
            return f"{entity_name} ({legal_form})"
        
        # Ajouter les fonctions à l'environnement Jinja2
        self.jinja_env.globals.update({
            'format_currency': format_currency,
            'format_date': format_date,
            'format_percentage': format_percentage,
            'format_number': format_number,
            'upper_first': upper_first,
            'legal_entity_format': legal_entity_format,
            'today': datetime.now(),
            'current_year': datetime.now().year
        })
    
    def render_template(self, template_content: str, variables: Dict[str, Any]) -> str:
        """Rend un template avec les variables fournies"""
        
        try:
            template = Template(template_content, environment=self.jinja_env)
            rendered = template.render(**variables)
            
            return rendered
            
        except TemplateError as e:
            logger.error(f"❌ Erreur rendu template: {e}")
            raise ValueError(f"Erreur de template: {e}")
        except Exception as e:
            logger.error(f"❌ Erreur inattendue rendu: {e}")
            raise
    
    def validate_variables(self, template_metadata: TemplateMetadata, variables: Dict[str, Any]) -> List[str]:
        """Valide les variables par rapport au template"""
        
        errors = []
        
        for var_def in template_metadata.variables:
            value = variables.get(var_def.name)
            
            # Vérifier variables requises
            if var_def.required and (value is None or value == ""):
                errors.append(f"Variable requise manquante: {var_def.name}")
                continue
            
            if value is not None:
                # Validation du type
                if var_def.type == "number" and not isinstance(value, (int, float)):
                    try:
                        float(value)
                    except (ValueError, TypeError):
                        errors.append(f"Variable {var_def.name} doit être un nombre")
                
                elif var_def.type == "date" and not isinstance(value, (datetime, str)):
                    errors.append(f"Variable {var_def.name} doit être une date")
                
                elif var_def.type == "boolean" and not isinstance(value, bool):
                    if str(value).lower() not in ['true', 'false', '1', '0']:
                        errors.append(f"Variable {var_def.name} doit être un booléen")
                
                # Validation pattern
                if var_def.validation_pattern and isinstance(value, str):
                    if not re.match(var_def.validation_pattern, value):
                        errors.append(f"Variable {var_def.name} ne respecte pas le format requis")
                
                # Validation choix
                if var_def.choices and value not in var_def.choices:
                    errors.append(f"Variable {var_def.name} doit être l'une des valeurs: {', '.join(var_def.choices)}")
        
        return errors


class PDFGenerator:
    """Générateur de documents PDF"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Configure les styles personnalisés"""
        
        # Style titre principal
        self.styles.add(ParagraphStyle(
            name='MainTitle',
            parent=self.styles['Title'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        ))
        
        # Style sous-titre
        self.styles.add(ParagraphStyle(
            name='Subtitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkgrey
        ))
        
        # Style signature
        self.styles.add(ParagraphStyle(
            name='Signature',
            parent=self.styles['Normal'],
            fontSize=10,
            alignment=TA_RIGHT,
            spaceBefore=30,
            textColor=colors.grey
        ))
        
        # Style clauses légales
        self.styles.add(ParagraphStyle(
            name='LegalClause',
            parent=self.styles['Normal'],
            fontSize=9,
            spaceBefore=6,
            spaceAfter=6,
            leftIndent=20,
            rightIndent=20
        ))
    
    def generate_pdf(self, content: str, metadata: Dict[str, Any] = None) -> bytes:
        """Génère un PDF à partir du contenu HTML/texte"""
        
        try:
            # Créer document temporaire
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                doc = SimpleDocTemplate(
                    tmp_file.name,
                    pagesize=A4,
                    rightMargin=0.75*inch,
                    leftMargin=0.75*inch,
                    topMargin=1*inch,
                    bottomMargin=1*inch
                )
                
                # Contenu du document
                story = []
                
                # Ajouter titre si présent
                if metadata and metadata.get('title'):
                    title = Paragraph(metadata['title'], self.styles['MainTitle'])
                    story.append(title)
                    story.append(Spacer(1, 12))
                
                # Parser le contenu (simplifié pour cette démo)
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        if line.startswith('# '):
                            # Titre
                            p = Paragraph(line[2:], self.styles['Heading1'])
                        elif line.startswith('## '):
                            # Sous-titre
                            p = Paragraph(line[3:], self.styles['Subtitle'])
                        elif line.startswith('**') and line.endswith('**'):
                            # Texte en gras
                            p = Paragraph(f"<b>{line[2:-2]}</b>", self.styles['Normal'])
                        else:
                            # Texte normal
                            p = Paragraph(line, self.styles['Normal'])
                        
                        story.append(p)
                        story.append(Spacer(1, 6))
                
                # Ajouter pied de page
                if metadata and metadata.get('generated_by'):
                    footer = Paragraph(
                        f"Document généré le {datetime.now().strftime('%d/%m/%Y')} par {metadata['generated_by']}",
                        self.styles['Signature']
                    )
                    story.append(Spacer(1, 20))
                    story.append(footer)
                
                # Construire le PDF
                doc.build(story)
                
                # Lire le fichier généré
                with open(tmp_file.name, 'rb') as f:
                    pdf_data = f.read()
                
                # Nettoyer
                os.unlink(tmp_file.name)
                
                return pdf_data
                
        except Exception as e:
            logger.error(f"❌ Erreur génération PDF: {e}")
            raise


class DOCXGenerator:
    """Générateur de documents DOCX"""
    
    def generate_docx(self, content: str, metadata: Dict[str, Any] = None) -> bytes:
        """Génère un DOCX à partir du contenu"""
        
        try:
            # Créer document
            doc = Document()
            
            # Ajouter titre si présent
            if metadata and metadata.get('title'):
                title = doc.add_heading(metadata['title'], level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parser le contenu (simplifié)
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    if line.startswith('# '):
                        # Titre
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        # Sous-titre
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('**') and line.endswith('**'):
                        # Texte en gras
                        p = doc.add_paragraph()
                        run = p.add_run(line[2:-2])
                        run.bold = True
                    else:
                        # Texte normal
                        doc.add_paragraph(line)
            
            # Ajouter pied de page
            if metadata and metadata.get('generated_by'):
                doc.add_page_break()
                footer = doc.add_paragraph(
                    f"Document généré le {datetime.now().strftime('%d/%m/%Y')} par {metadata['generated_by']}"
                )
                footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Sauvegarder en mémoire
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                doc.save(tmp_file.name)
                
                with open(tmp_file.name, 'rb') as f:
                    docx_data = f.read()
                
                os.unlink(tmp_file.name)
                
                return docx_data
                
        except Exception as e:
            logger.error(f"❌ Erreur génération DOCX: {e}")
            raise


class DocumentTemplateManager:
    """Gestionnaire principal des templates de documents"""
    
    def __init__(self, templates_store_path: str = "templates_metadata.json"):
        self.templates_store_path = templates_store_path
        self.templates_store: Dict[str, TemplateMetadata] = {}
        
        # Moteurs de génération
        self.template_engine = TemplateEngine_()
        self.pdf_generator = PDFGenerator()
        self.docx_generator = DOCXGenerator()
        
        # Cache
        self.cache = get_cache_manager()
        
        # Statistiques
        self.generation_stats = {
            "total_generations": 0,
            "generations_by_type": {},
            "generations_by_format": {},
            "average_generation_time": 0.0
        }
    
    async def initialize(self):
        """Initialise le gestionnaire de templates"""
        try:
            logger.info("🚀 Initialisation du gestionnaire de templates...")
            
            # Charger templates existants
            await self._load_templates_store()
            
            # Créer templates par défaut si aucun n'existe
            if not self.templates_store:
                await self._create_default_templates()
            
            logger.info("✅ Gestionnaire de templates initialisé")
            
        except Exception as e:
            logger.error(f"❌ Erreur initialisation templates: {e}")
            raise
    
    async def _load_templates_store(self):
        """Charge le store de templates"""
        try:
            if os.path.exists(self.templates_store_path):
                with open(self.templates_store_path, 'r') as f:
                    data = json.load(f)
                
                for template_id, template_data in data.items():
                    self.templates_store[template_id] = TemplateMetadata.from_dict(template_data)
                
                logger.info(f"📂 {len(self.templates_store)} templates chargés")
        
        except Exception as e:
            logger.warning(f"⚠️ Erreur chargement templates: {e}")
    
    async def _save_templates_store(self):
        """Sauvegarde le store de templates"""
        try:
            data = {
                template_id: template_metadata.to_dict()
                for template_id, template_metadata in self.templates_store.items()
            }
            
            with open(self.templates_store_path, 'w') as f:
                json.dump(data, f, indent=2)
        
        except Exception as e:
            logger.error(f"❌ Erreur sauvegarde templates: {e}")
    
    async def _create_default_templates(self):
        """Crée les templates par défaut"""
        
        # Template NDA
        nda_template = TemplateMetadata(
            template_id=str(uuid.uuid4()),
            name="Accord de Confidentialité Standard",
            description="Template d'accord de non-divulgation pour opérations M&A",
            template_type=TemplateType.NDA,
            variables=[
                TemplateVariable("partie_1_nom", "string", "Nom de la première partie", required=True),
                TemplateVariable("partie_1_adresse", "string", "Adresse de la première partie", required=True),
                TemplateVariable("partie_2_nom", "string", "Nom de la seconde partie", required=True),
                TemplateVariable("partie_2_adresse", "string", "Adresse de la seconde partie", required=True),
                TemplateVariable("objet_transaction", "string", "Objet de la transaction", required=True),
                TemplateVariable("duree_confidentialite", "number", "Durée de confidentialité en années", default_value=5),
                TemplateVariable("juridiction", "string", "Juridiction applicable", default_value="France"),
                TemplateVariable("date_signature", "date", "Date de signature", required=True)
            ],
            tags=["nda", "confidentialité", "m&a", "légal"]
        )
        
        # Template LOI
        loi_template = TemplateMetadata(
            template_id=str(uuid.uuid4()),
            name="Lettre d'Intention",
            description="Template de lettre d'intention pour acquisition",
            template_type=TemplateType.LOI,
            variables=[
                TemplateVariable("acquereur_nom", "string", "Nom de l'acquéreur", required=True),
                TemplateVariable("cible_nom", "string", "Nom de l'entreprise cible", required=True),
                TemplateVariable("prix_propose", "number", "Prix proposé", required=True),
                TemplateVariable("devise", "string", "Devise", default_value="EUR"),
                TemplateVariable("modalites_paiement", "string", "Modalités de paiement"),
                TemplateVariable("conditions_suspensives", "list", "Conditions suspensives"),
                TemplateVariable("echeance_due_diligence", "date", "Échéance due diligence"),
                TemplateVariable("echeance_offre", "date", "Échéance de l'offre", required=True)
            ],
            tags=["loi", "intention", "acquisition", "m&a"]
        )
        
        # Template Due Diligence
        dd_template = TemplateMetadata(
            template_id=str(uuid.uuid4()),
            name="Rapport de Due Diligence",
            description="Template de rapport de due diligence",
            template_type=TemplateType.DUE_DILIGENCE,
            variables=[
                TemplateVariable("entreprise_nom", "string", "Nom de l'entreprise", required=True),
                TemplateVariable("secteur_activite", "string", "Secteur d'activité", required=True),
                TemplateVariable("chiffre_affaires", "number", "Chiffre d'affaires"),
                TemplateVariable("effectifs", "number", "Nombre d'employés"),
                TemplateVariable("points_forts", "list", "Points forts identifiés"),
                TemplateVariable("points_faibles", "list", "Points faibles identifiés"),
                TemplateVariable("risques_majeurs", "list", "Risques majeurs"),
                TemplateVariable("recommandations", "list", "Recommandations"),
                TemplateVariable("auditeur_nom", "string", "Nom de l'auditeur", required=True)
            ],
            tags=["due-diligence", "audit", "analyse", "m&a"]
        )
        
        # Sauvegarder templates
        self.templates_store[nda_template.template_id] = nda_template
        self.templates_store[loi_template.template_id] = loi_template
        self.templates_store[dd_template.template_id] = dd_template
        
        await self._save_templates_store()
        
        # Créer fichiers templates
        await self._create_template_files()
        
        logger.info("📝 Templates par défaut créés")
    
    async def _create_template_files(self):
        """Crée les fichiers templates par défaut"""
        
        # Template NDA
        nda_content = """
# ACCORD DE CONFIDENTIALITÉ

**ENTRE :**

**{{ partie_1_nom }}**
{{ partie_1_adresse }}
Ci-après dénommée "la Partie Divulgatrice"

**ET :**

**{{ partie_2_nom }}**
{{ partie_2_adresse }}
Ci-après dénommée "la Partie Bénéficiaire"

## ARTICLE 1 - OBJET

Dans le cadre de {{ objet_transaction }}, la Partie Divulgatrice sera amenée à communiquer à la Partie Bénéficiaire des informations confidentielles.

## ARTICLE 2 - DÉFINITION DES INFORMATIONS CONFIDENTIELLES

Sont considérées comme confidentielles toutes informations, documents, données, procédés, méthodes, savoir-faire, de quelque nature que ce soit, communiqués par la Partie Divulgatrice.

## ARTICLE 3 - OBLIGATIONS DE CONFIDENTIALITÉ

La Partie Bénéficiaire s'engage à :
- Maintenir strictement confidentielles toutes les informations reçues
- Ne pas divulguer ces informations à des tiers sans autorisation écrite préalable
- Utiliser ces informations uniquement dans le cadre défini

## ARTICLE 4 - DURÉE

Le présent accord prend effet à compter du {{ format_date(date_signature) }} et demeure en vigueur pendant {{ duree_confidentialite }} années.

## ARTICLE 5 - DROIT APPLICABLE

Le présent accord est régi par le droit {{ juridiction }}.

**Fait le {{ format_date(today) }}**

{{ partie_1_nom }}                    {{ partie_2_nom }}
"""
        
        # Template LOI
        loi_content = """
# LETTRE D'INTENTION

**De :** {{ acquereur_nom }}
**À :** {{ cible_nom }}
**Date :** {{ format_date(today) }}

## OBJET : INTENTION D'ACQUISITION

Nous avons l'honneur de vous faire part de notre intention d'acquérir {{ cible_nom }}.

## PROPOSITION FINANCIÈRE

**Prix proposé :** {{ format_currency(prix_propose, devise) }}

**Modalités de paiement :** {{ modalites_paiement or "À définir lors des négociations" }}

## CONDITIONS DE L'OFFRE

### Conditions suspensives
{% if conditions_suspensives %}
{% for condition in conditions_suspensives %}
- {{ condition }}
{% endfor %}
{% else %}
- Due diligence satisfaisante
- Obtention des autorisations nécessaires
- Financement confirmé
{% endif %}

### Calendrier
- **Échéance due diligence :** {{ format_date(echeance_due_diligence) if echeance_due_diligence else "À définir" }}
- **Échéance de l'offre :** {{ format_date(echeance_offre) }}

## CONFIDENTIALITÉ

Cette lettre d'intention et toutes les négociations qui en découlent sont strictement confidentielles.

**{{ acquereur_nom }}**
*Le {{ format_date(today) }}*
"""
        
        # Template Due Diligence
        dd_content = """
# RAPPORT DE DUE DILIGENCE

**Entreprise analysée :** {{ entreprise_nom }}
**Secteur d'activité :** {{ secteur_activite }}
**Date du rapport :** {{ format_date(today) }}
**Auditeur :** {{ auditeur_nom }}

## RÉSUMÉ EXÉCUTIF

### Données clés
{% if chiffre_affaires %}
- **Chiffre d'affaires :** {{ format_currency(chiffre_affaires) }}
{% endif %}
{% if effectifs %}
- **Effectifs :** {{ format_number(effectifs) }} employés
{% endif %}

## ANALYSE DÉTAILLÉE

### Points forts identifiés
{% if points_forts %}
{% for point in points_forts %}
- {{ point }}
{% endfor %}
{% else %}
- Analyse en cours
{% endif %}

### Points faibles identifiés
{% if points_faibles %}
{% for point in points_faibles %}
- {{ point }}
{% endfor %}
{% else %}
- Analyse en cours
{% endif %}

### Risques majeurs
{% if risques_majeurs %}
{% for risque in risques_majeurs %}
- **RISQUE :** {{ risque }}
{% endfor %}
{% else %}
- Analyse en cours
{% endif %}

## RECOMMANDATIONS

{% if recommandations %}
{% for recommandation in recommandations %}
{{ loop.index }}. {{ recommandation }}
{% endfor %}
{% else %}
Les recommandations seront fournies après analyse complète.
{% endif %}

## CONCLUSION

Ce rapport de due diligence a été établi sur la base des informations disponibles au {{ format_date(today) }}.

---
*Rapport établi par {{ auditeur_nom }}*
*{{ format_date(today) }}*
"""
        
        # Sauvegarder les fichiers
        templates_content = {
            "nda.j2": nda_content,
            "loi.j2": loi_content,
            "due_diligence.j2": dd_content
        }
        
        for filename, content in templates_content.items():
            file_path = self.template_engine.templates_dir / filename
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content.strip())
    
    async def generate_document(self, request: GenerationRequest) -> GenerationResult:
        """Génère un document à partir d'un template"""
        
        start_time = datetime.now()
        
        try:
            # Récupérer métadonnées du template
            template_metadata = self.templates_store.get(request.template_id)
            if not template_metadata:
                raise ValueError(f"Template non trouvé: {request.template_id}")
            
            # Valider variables
            validation_errors = self.template_engine.validate_variables(template_metadata, request.variables)
            if validation_errors:
                return GenerationResult(
                    success=False,
                    errors=validation_errors
                )
            
            # Enrichir variables avec valeurs par défaut
            enriched_variables = {}
            for var_def in template_metadata.variables:
                if var_def.name in request.variables:
                    enriched_variables[var_def.name] = request.variables[var_def.name]
                elif var_def.default_value is not None:
                    enriched_variables[var_def.name] = var_def.default_value
            
            # Ajouter variables système
            enriched_variables.update({
                'generated_by': request.generated_by,
                'generation_date': datetime.now(),
                'template_name': template_metadata.name,
                'template_version': template_metadata.version
            })
            
            # Charger et rendre template
            template_filename = self._get_template_filename(template_metadata.template_type)
            template_path = self.template_engine.templates_dir / template_filename
            
            if not template_path.exists():
                raise FileNotFoundError(f"Fichier template non trouvé: {template_filename}")
            
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()
            
            # Rendre le template
            rendered_content = self.template_engine.render_template(template_content, enriched_variables)
            
            # Générer dans le format demandé
            file_data, mime_type = await self._generate_output(
                rendered_content, 
                request.output_format,
                {
                    'title': template_metadata.name,
                    'generated_by': request.generated_by
                }
            )
            
            # Calculer statistiques
            generation_time = (datetime.now() - start_time).total_seconds()
            word_count = len(rendered_content.split())
            
            # Générer nom de fichier
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            template_name_safe = re.sub(r'[^\w\-_.]', '_', template_metadata.name.lower())
            filename = f"{template_name_safe}_{timestamp}.{request.output_format.value}"
            
            # Mettre à jour statistiques
            template_metadata.usage_count += 1
            self.generation_stats["total_generations"] += 1
            type_count = self.generation_stats["generations_by_type"].get(template_metadata.template_type.value, 0)
            self.generation_stats["generations_by_type"][template_metadata.template_type.value] = type_count + 1
            format_count = self.generation_stats["generations_by_format"].get(request.output_format.value, 0)
            self.generation_stats["generations_by_format"][request.output_format.value] = format_count + 1
            
            # Sauvegarder templates
            await self._save_templates_store()
            
            result = GenerationResult(
                success=True,
                file_data=file_data,
                filename=filename,
                mime_type=mime_type,
                template_used=template_metadata.name,
                generation_time=generation_time,
                pages_generated=1,  # Estimation
                word_count=word_count
            )
            
            logger.info(f"📄 Document généré: {template_metadata.name} -> {request.output_format.value}")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Erreur génération document: {e}")
            return GenerationResult(
                success=False,
                errors=[str(e)],
                generation_time=(datetime.now() - start_time).total_seconds()
            )
    
    async def _generate_output(self, content: str, output_format: OutputFormat, metadata: Dict[str, Any]) -> Tuple[bytes, str]:
        """Génère la sortie dans le format demandé"""
        
        if output_format == OutputFormat.PDF:
            file_data = self.pdf_generator.generate_pdf(content, metadata)
            return file_data, "application/pdf"
        
        elif output_format == OutputFormat.DOCX:
            file_data = self.docx_generator.generate_docx(content, metadata)
            return file_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        elif output_format == OutputFormat.HTML:
            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{metadata.get('title', 'Document')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #2c3e50; }}
        h2 {{ color: #34495e; }}
        .footer {{ margin-top: 50px; font-size: 12px; color: #7f8c8d; text-align: right; }}
    </style>
</head>
<body>
{content.replace(chr(10), '<br>')}
<div class="footer">
    Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}
</div>
</body>
</html>
"""
            return html_content.encode('utf-8'), "text/html"
        
        elif output_format == OutputFormat.TXT:
            return content.encode('utf-8'), "text/plain"
        
        elif output_format == OutputFormat.JSON:
            json_data = {
                "content": content,
                "metadata": metadata,
                "generated_at": datetime.now().isoformat()
            }
            return json.dumps(json_data, indent=2).encode('utf-8'), "application/json"
        
        else:
            raise ValueError(f"Format de sortie non supporté: {output_format}")
    
    def _get_template_filename(self, template_type: TemplateType) -> str:
        """Retourne le nom de fichier pour un type de template"""
        
        filename_mapping = {
            TemplateType.NDA: "nda.j2",
            TemplateType.LOI: "loi.j2",
            TemplateType.DUE_DILIGENCE: "due_diligence.j2",
            TemplateType.TERM_SHEET: "term_sheet.j2",
            TemplateType.FINANCIAL_REPORT: "financial_report.j2",
            TemplateType.LEGAL_OPINION: "legal_opinion.j2",
            TemplateType.VALUATION_REPORT: "valuation_report.j2"
        }
        
        return filename_mapping.get(template_type, "default.j2")
    
    def list_templates(self, template_type: TemplateType = None) -> List[TemplateMetadata]:
        """Liste les templates disponibles"""
        
        templates = list(self.templates_store.values())
        
        if template_type:
            templates = [t for t in templates if t.template_type == template_type]
        
        return sorted(templates, key=lambda t: t.name)
    
    def get_template(self, template_id: str) -> Optional[TemplateMetadata]:
        """Récupère un template par son ID"""
        return self.templates_store.get(template_id)
    
    def get_generation_statistics(self) -> Dict[str, Any]:
        """Retourne les statistiques de génération"""
        return {
            **self.generation_stats,
            "total_templates": len(self.templates_store),
            "templates_by_type": {
                template_type.value: len([t for t in self.templates_store.values() 
                                        if t.template_type == template_type])
                for template_type in TemplateType
            }
        }


# Instance globale
_document_template_manager: Optional[DocumentTemplateManager] = None


async def get_document_template_manager() -> DocumentTemplateManager:
    """Factory pour obtenir le gestionnaire de templates"""
    global _document_template_manager
    
    if _document_template_manager is None:
        _document_template_manager = DocumentTemplateManager()
        await _document_template_manager.initialize()
    
    return _document_template_manager